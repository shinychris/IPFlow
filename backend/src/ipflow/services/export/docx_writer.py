"""DOCX 导出工具.

为软著/专利/商标材料导出提供 ``.docx`` 生成能力。基于 ``python-docx``，
将结构化章节（标题 + 段落 + 列表）渲染为 Word 文档。

若 ``python-docx`` 未安装，``render_docx`` 抛出 ``RuntimeError``，
调用方应捕获后降级为 TXT 输出。
"""

from __future__ import annotations

import io
import logging
from typing import Any, Optional, Union

logger = logging.getLogger(__name__)


def _import_docx():
    """惰性导入 docx，缺失时抛 RuntimeError."""
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
        return Document, Pt
    except ImportError as e:
        raise RuntimeError("python-docx 未安装，无法生成 DOCX") from e


def render_docx(
    title: str,
    sections: list[dict[str, Any]],
    *,
    subtitle: Optional[str] = None,
) -> bytes:
    """渲染 DOCX 文档。

    Args:
        title: 文档主标题（居中、加粗大字号）。
        sections: 章节列表，每个章节为 dict，支持字段：

            - ``heading`` (str): 章节标题（Heading 1）
            - ``paragraphs`` (list[str]): 普通段落
            - ``list_items`` (list[str]): 项目符号列表
            - ``key_values`` (list[tuple[str,str]]): 「字段：值」键值对表格
            - ``html`` (str): 简易 HTML 片段（仅解析 ``<h1>/<h2>/<p>/<li>``）

        subtitle: 可选副标题（标题下方居中显示）。

    Returns:
        DOCX 文件的字节内容。

    Raises:
        RuntimeError: ``python-docx`` 未安装。
    """
    Document, _ = _import_docx()  # noqa: F811
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

    doc = Document()

    # 主标题
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sub.runs:
            run.italic = True

    for section in sections:
        heading = section.get("heading")
        if heading:
            doc.add_heading(heading, level=section.get("level", 1))

        # 键值对表格
        key_values = section.get("key_values")
        if key_values:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "项目"
            hdr[1].text = "内容"
            for k, v in key_values:
                row = table.add_row().cells
                row[0].text = str(k)
                row[1].text = "" if v is None else str(v)
            doc.add_paragraph("")

        # 普通段落
        for para in section.get("paragraphs", []):
            doc.add_paragraph(str(para))

        # 项目符号列表
        for item in section.get("list_items", []):
            doc.add_paragraph(str(item), style="List Bullet")

        # 简易 HTML 片段
        html = section.get("html")
        if html:
            _render_html_to_docx(doc, html)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_html_to_docx(doc, html: str) -> None:
    """将极简 HTML 片段（h1/h2/p/li）追加到文档。"""
    import re

    # 按 HTML 标签切分
    tokens = re.findall(r"<(h[12]|p|li)[^>]*>(.*?)</\1>", html, flags=re.DOTALL)
    if not tokens:
        # 无标签时按换行拆为段落
        for line in re.sub(r"<[^>]+>", "", html).split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line)
        return
    for tag, content in tokens:
        text = re.sub(r"<[^>]+>", "", content).strip()
        if not text:
            continue
        if tag == "h1":
            doc.add_heading(text, level=1)
        elif tag == "h2":
            doc.add_heading(text, level=2)
        elif tag == "li":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)


def docx_available() -> bool:
    """检查 python-docx 是否可用。"""
    try:
        _import_docx()
        return True
    except RuntimeError:
        return False
