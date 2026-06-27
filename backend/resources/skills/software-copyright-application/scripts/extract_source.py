#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
软著源代码提取工具（生成 .docx）

从项目中提取并生成符合中国版权保护中心 docx 规范的源代码文档：
  - 纸张：A4（210×297mm）
  - 页边距：常规（上下 1"，左右 0.8-1.25"）
  - 页眉：软件全称 + 版本号（居中，所有页面显示）
  - 页码：底部居中，连续编排
  - 代码字体：Consolas / Courier New（等宽）
  - 字号：五号（10.5pt）
  - 行距：单倍
  - 每页代码量：≥ 50 行（不含空行）
  - 文件标识：每个源文件前注明「代码文件: <相对路径>」
  - 注释率：<5%（由 ai_trace_analyzer.py 核查）

依赖：python-docx
    pip install python-docx

使用方法:
    python extract_source.py --project-path /path/to/project \
        --software-name "XX管理系统" \
        --version "V1.0" \
        --output-dir ./output
"""

import os
import re
import argparse
from pathlib import Path
from datetime import datetime


# 默认支持的代码文件扩展名
DEFAULT_EXTENSIONS = {
    '.java', '.py', '.js', '.ts', '.jsx', '.tsx',
    '.c', '.cpp', '.h', '.hpp', '.cs', '.go',
    '.php', '.rb', '.swift', '.kt', '.scala',
    '.vue', '.html', '.css', '.scss', '.less',
    '.sql', '.xml', '.json', '.yaml', '.yml',
    '.sh', '.bat', '.ps1', '.r', '.m', '.mm'
}

# 默认排除的目录
DEFAULT_EXCLUDE_DIRS = {
    '.git', '.svn', '.hg', 'node_modules', 'vendor',
    'target', 'build', 'dist', 'out', '.idea', '.vscode',
    '__pycache__', '.gradle', 'bin', 'obj', 'release',
    'debug', 'logs', 'temp', 'tmp', 'uploads', 'assets',
    'public', 'static', 'test', 'tests', 'spec', 'mock'
}

# 默认排除的文件
DEFAULT_EXCLUDE_FILES = {
    '.gitignore', '.gitattributes', '.editorconfig',
    '.dockerignore', 'LICENSE', 'README', 'CHANGELOG',
    'Makefile', 'CMakeLists.txt', 'package-lock.json',
    'yarn.lock', 'pnpm-lock.yaml', 'Cargo.lock',
    'requirements.txt', 'Gemfile.lock', 'composer.lock'
}

# 注释行匹配模式（用于估算注释率，<5% 才达标）
COMMENT_PATTERNS = [
    re.compile(r'^\s*#'),                       # Python/Ruby/Shell
    re.compile(r'^\s*//'),                      # C/Java/JS/Go/Rust
    re.compile(r'^\s*/\*'),                     # C 风格块注释开始
    re.compile(r'^\s*\*'),                      # 块注释续行
    re.compile(r'^\s*<!--'),                    # HTML/XML
    re.compile(r'^\s*--'),                      # SQL/Lua
    re.compile(r'^\s*"""'),                     # Python 文档字符串
    re.compile(r"^\s*'''"),
]


def get_all_code_files(project_path, extensions=None, exclude_dirs=None, exclude_files=None):
    """获取项目中所有代码文件"""
    if extensions is None:
        extensions = DEFAULT_EXTENSIONS
    if exclude_dirs is None:
        exclude_dirs = DEFAULT_EXCLUDE_DIRS
    if exclude_files is None:
        exclude_files = DEFAULT_EXCLUDE_FILES

    code_files = []
    project_path = Path(project_path).resolve()

    for root, dirs, files in os.walk(project_path):
        # 排除指定目录
        dirs[:] = [d for d in dirs if d not in exclude_dirs and not d.startswith('.')]

        for file in files:
            # 跳过隐藏文件
            if file.startswith('.'):
                continue
            # 跳过排除的文件
            if file in exclude_files:
                continue
            # 检查扩展名
            file_path = Path(root) / file
            if file_path.suffix.lower() in extensions:
                code_files.append(file_path)

    return sorted(code_files)


def count_lines(file_path):
    """统计文件行数"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return len(f.readlines())
    except Exception:
        return 0


def remove_blank_lines(content):
    """移除空行"""
    lines = content.split('\n')
    non_blank_lines = [line for line in lines if line.strip()]
    return '\n'.join(non_blank_lines)


def remove_sensitive_info(content):
    """移除敏感信息（简单处理）"""
    patterns = [
        (r'(password|passwd|pwd)\s*[=:]\s*["\'][^"\']+["\']', r'\1="***"'),
        (r'(secret|api[_-]?key|token)\s*[=:]\s*["\'][^"\']+["\']', r'\1="***"'),
        (r'(jdbc:|mysql://|postgresql://)[^\s"\']+', r'\1***'),
    ]

    for pattern, repl in patterns:
        content = re.sub(pattern, repl, content, flags=re.IGNORECASE)

    return content


def read_file_content(file_path, remove_blank=True, desensitize=True):
    """读取文件内容并进行处理"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        if remove_blank:
            content = remove_blank_lines(content)

        if desensitize:
            content = remove_sensitive_info(content)

        return content
    except Exception as e:
        return f"// Error reading file: {e}\n"


def collect_code_content(code_files, project_root=None, max_lines=None):
    """
    收集代码内容，每个源文件前注明「代码文件: <相对路径>」

    Args:
        code_files: 代码文件路径列表
        project_root: 项目根目录（用于计算相对路径）
        max_lines: 最大收集行数，None表示不限制

    Returns:
        str: 收集的代码内容
    """
    all_content = []
    total_lines = 0

    if project_root is None and code_files:
        project_root = code_files[0].parent

    for file_path in code_files:
        # 计算相对路径作为文件标识
        try:
            rel_path = file_path.relative_to(project_root) if project_root else file_path.name
            rel_path_str = str(rel_path).replace('\\', '/')
        except ValueError:
            rel_path_str = file_path.name

        # 文件标识（软著规范要求）
        separator = f"\n// 代码文件: {rel_path_str}\n\n"

        content = read_file_content(file_path)
        file_lines = content.count('\n') + 1

        if max_lines and total_lines + file_lines > max_lines:
            # 只取需要的部分
            remaining = max_lines - total_lines
            lines = content.split('\n')[:remaining]
            content = '\n'.join(lines)
            all_content.append(separator + content)
            break

        all_content.append(separator + content)
        total_lines += file_lines

    return '\n'.join(all_content)


def split_into_pages(content, lines_per_page=50):
    """
    将代码内容按指定行数分页（每页 ≥50 行）

    Args:
        content: 代码内容
        lines_per_page: 每页行数（默认 50，符合软著规范）

    Returns:
        list: 每页内容的列表
    """
    lines = content.split('\n')
    pages = []

    for i in range(0, len(lines), lines_per_page):
        page_lines = lines[i:i + lines_per_page]
        pages.append('\n'.join(page_lines))

    return pages


def _set_cell_margins(section, top=1.0, bottom=1.0, left=1.0, right=1.0):
    """设置页边距（英寸）"""
    from docx.shared import Inches
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def _set_a4_paper(section):
    """设置 A4 纸张（210×297mm）"""
    from docx.shared import Mm
    section.page_width = Mm(210)
    section.page_height = Mm(297)


def _set_code_font(run):
    """设置代码字体为 Consolas / Courier New，五号（10.5pt）"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    run.font.name = 'Consolas'
    # 兼容 Courier New（若环境无 Consolas）
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    rFonts.set(qn('w:cs'), 'Courier New')
    run.font.size = Pt(10.5)


def _add_page_number(paragraph):
    """在段落中插入页码域（底部居中，连续编排）"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    # 插入 PAGE 域
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _setup_header_footer(section, software_name, version):
    """
    设置页眉（软件全称+版本号，居中）和页脚（页码，底部居中）
    符合软著 docx 规范
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    # 页眉：软件全称 + 版本号，居中，小五号
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.text = f"{software_name} {version}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header_para.runs:
        run.font.size = Pt(9)
        run.font.name = '宋体'

    # 页脚：页码，底部居中，连续编排
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(footer_para)
    for run in footer_para.runs:
        run.font.size = Pt(9)


def generate_document(pages, software_name, version, output_path, pages_per_doc=30, lines_per_page=50):
    """
    生成符合软著 docx 规范的源代码文档

    规范要点：
      - A4 纸张
      - 常规页边距
      - 页眉：软件全称+版本号（居中）
      - 页脚：页码（底部居中，连续）
      - 代码字体 Consolas / Courier New，五号（10.5pt）
      - 单倍行距
      - 每页 ≥50 行

    Args:
        pages: 页面内容列表
        software_name: 软件全称
        version: 版本号
        output_path: 输出文件路径（.docx）
        pages_per_doc: 每个文档的页数
        lines_per_page: 每页行数
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("错误：必须安装 python-docx 才能生成 .docx")
        print("请运行: pip install python-docx")
        return False

    doc = Document()

    # 页面设置：A4 + 常规页边距
    section = doc.sections[0]
    _set_a4_paper(section)
    _set_cell_margins(section, top=1.0, bottom=1.0, left=1.0, right=1.0)

    # 页眉页脚
    _setup_header_footer(section, software_name, version)

    # 添加代码内容
    selected_pages = pages[:pages_per_doc]
    for i, page_content in enumerate(selected_pages):
        if i > 0:
            doc.add_page_break()

        # 按行添加，保证每页行数与字体可控
        lines = page_content.split('\n')
        # 不足 lines_per_page 行不补空行（规范要求 ≥50 行非空）
        for line in lines:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line if line else ' ')
            _set_code_font(run)
            # 单倍行距，段后无空隙
            pf = paragraph.paragraph_format
            pf.line_spacing = 1.0
            pf.space_after = Pt(0)
            pf.space_before = Pt(0)

    doc.save(output_path)
    return True


def estimate_comment_ratio(content):
    """估算注释率（注释行 / 非空行），用于自查 <5%"""
    lines = [ln for ln in content.split('\n') if ln.strip()]
    if not lines:
        return 0.0
    comment_lines = sum(1 for ln in lines if any(p.match(ln) for p in COMMENT_PATTERNS))
    return comment_lines / len(lines)


def main():
    parser = argparse.ArgumentParser(
        description='软著源代码提取工具（生成 .docx，符合 A4/Consolas/页眉规范）',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 基本用法
  python extract_source.py --project-path ./my-project --software-name "XX系统" --version "V1.0"

  # 自定义输出目录
  python extract_source.py -p ./my-project -n "XX系统" -v "V1.0" -o ./output

  # 自定义文件扩展名
  python extract_source.py -p ./my-project -n "XX系统" -v "V1.0" --ext .java,.xml,.properties

  # 指定自定义的前后30页代码
  python extract_source.py -p ./my-project -n "XX系统" -v "V1.0" \\
      --front-pages ./custom/front.txt --back-pages ./custom/back.txt

注意：本工具仅输出 .docx 格式（不再支持 txt），需先 pip install python-docx
        """
    )

    parser.add_argument('-p', '--project-path', required=True,
                        help='项目根目录路径')
    parser.add_argument('-n', '--software-name', required=True,
                        help='软件全称（用于页眉）')
    parser.add_argument('-v', '--version', required=True,
                        help='版本号（用于页眉，如 V1.0）')
    parser.add_argument('-o', '--output-dir', default='./软著源代码',
                        help='输出目录（默认: ./软著源代码）')
    parser.add_argument('--ext', type=str,
                        help='包含的文件扩展名，逗号分隔（如: .java,.xml）')
    parser.add_argument('--exclude-dirs', type=str,
                        help='排除的目录，逗号分隔')
    parser.add_argument('--exclude-files', type=str,
                        help='排除的文件，逗号分隔')
    parser.add_argument('--lines-per-page', type=int, default=50,
                        help='每页行数（默认: 50，符合软著规范）')
    parser.add_argument('--front-pages', type=str,
                        help='自定义前30页代码文件路径')
    parser.add_argument('--back-pages', type=str,
                        help='自定义后30页代码文件路径')

    args = parser.parse_args()

    # 检查 python-docx 是否安装
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        print("错误：必须安装 python-docx 才能生成 .docx 文档")
        print("请运行: pip install python-docx")
        return

    # 解析扩展名
    if args.ext:
        extensions = set(ext.strip() for ext in args.ext.split(','))
    else:
        extensions = DEFAULT_EXTENSIONS

    # 解析排除项
    exclude_dirs = DEFAULT_EXCLUDE_DIRS.copy()
    if args.exclude_dirs:
        exclude_dirs.update(d.strip() for d in args.exclude_dirs.split(','))

    exclude_files = DEFAULT_EXCLUDE_FILES.copy()
    if args.exclude_files:
        exclude_files.update(f.strip() for f in args.exclude_files.split(','))

    # 创建输出目录
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"正在扫描项目: {args.project_path}")
    print(f"文件扩展名: {extensions}")

    # 获取所有代码文件
    code_files = get_all_code_files(
        args.project_path,
        extensions=extensions,
        exclude_dirs=exclude_dirs,
        exclude_files=exclude_files
    )

    print(f"找到 {len(code_files)} 个代码文件")

    if not code_files:
        print("错误：未找到任何代码文件，请检查项目路径和扩展名设置")
        return

    # 统计总代码行数
    total_lines = sum(count_lines(f) for f in code_files)
    print(f"总代码行数: {total_lines}")

    project_root = Path(args.project_path).resolve()

    # 生成前30页
    print("\n正在生成前30页源代码...")
    if args.front_pages and Path(args.front_pages).exists():
        # 使用自定义前30页
        with open(args.front_pages, 'r', encoding='utf-8') as f:
            front_content = f.read()
        front_pages = split_into_pages(front_content, args.lines_per_page)
        print(f"使用自定义前30页: {args.front_pages}")
    else:
        # 从项目中提取
        # 优先选择：配置文件 -> 入口文件 -> 核心模块
        priority_files = []
        other_files = []

        config_patterns = ['package.json', 'pom.xml', 'build.gradle', 'requirements.txt',
                           'Cargo.toml', 'go.mod', 'composer.json', 'Gemfile']
        main_patterns = ['main', 'app', 'index', 'application', 'startup']

        for f in code_files:
            fname = f.name.lower()
            if any(cp in fname for cp in config_patterns):
                priority_files.append(f)
            elif any(mp in fname for mp in main_patterns):
                priority_files.append(f)
            else:
                other_files.append(f)

        # 合并文件列表（配置文件优先，然后按字母顺序）
        sorted_files = priority_files + other_files
        front_content = collect_code_content(sorted_files, project_root=project_root,
                                              max_lines=30 * args.lines_per_page)
        front_pages = split_into_pages(front_content, args.lines_per_page)

    # 生成后30页
    print("正在生成后30页源代码...")
    if args.back_pages and Path(args.back_pages).exists():
        # 使用自定义后30页
        with open(args.back_pages, 'r', encoding='utf-8') as f:
            back_content = f.read()
        back_pages = split_into_pages(back_content, args.lines_per_page)
        print(f"使用自定义后30页: {args.back_pages}")
    else:
        # 从项目中倒序提取
        reversed_files = list(reversed(code_files))
        back_content = collect_code_content(reversed_files, project_root=project_root,
                                             max_lines=30 * args.lines_per_page)
        back_pages = split_into_pages(back_content, args.lines_per_page)

    # 生成 .docx 文档
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    front_path = output_dir / f"源代码_前30页_{timestamp}.docx"
    back_path = output_dir / f"源代码_后30页_{timestamp}.docx"

    print(f"\n生成 .docx 文档（A4 / Consolas 五号 / 页眉含软件名）...")
    ok_front = generate_document(front_pages, args.software_name, args.version,
                                  front_path, 30, args.lines_per_page)
    ok_back = generate_document(back_pages, args.software_name, args.version,
                                 back_path, 30, args.lines_per_page)

    if ok_front:
        print(f"前30页: {front_path}")
    if ok_back:
        print(f"后30页: {back_path}")

    # 注释率自查（软著硬性指标 <5%）
    front_ratio = estimate_comment_ratio(front_content)
    back_ratio = estimate_comment_ratio(back_content)
    print(f"\n⚠️ 注释率自查（目标 <5%）:")
    print(f"  前30页注释率: {front_ratio:.1%} {'✅' if front_ratio < 0.05 else '❌ 超标，需削减注释'}")
    print(f"  后30页注释率: {back_ratio:.1%} {'✅' if back_ratio < 0.05 else '❌ 超标，需削减注释'}")
    if front_ratio >= 0.05 or back_ratio >= 0.05:
        print("  建议：运行 python scripts/ai_trace_analyzer.py --type code --path <文件> 获取详细改造建议")

    # 生成统计信息
    info_path = output_dir / f"统计信息_{timestamp}.txt"
    with open(info_path, 'w', encoding='utf-8') as f:
        f.write(f"软件名称: {args.software_name}\n")
        f.write(f"版本号: {args.version}\n")
        f.write(f"项目路径: {args.project_path}\n")
        f.write(f"扫描时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"\n统计信息:\n")
        f.write(f"  代码文件总数: {len(code_files)}\n")
        f.write(f"  总代码行数: {total_lines}\n")
        f.write(f"  平均每文件行数: {total_lines // len(code_files) if code_files else 0}\n")
        f.write(f"  前30页注释率: {front_ratio:.1%}（目标 <5%）\n")
        f.write(f"  后30页注释率: {back_ratio:.1%}（目标 <5%）\n")
        f.write(f"\n文件扩展名分布:\n")

        ext_count = {}
        for f2 in code_files:
            ext = f2.suffix.lower()
            ext_count[ext] = ext_count.get(ext, 0) + 1

        for ext, count in sorted(ext_count.items(), key=lambda x: -x[1]):
            f.write(f"  {ext}: {count} 个文件\n")

        f.write(f"\n输出格式: .docx（A4 / Consolas 五号 / 单倍行距 / 页眉软件全称+版本号 / 页脚页码居中）\n")
        f.write(f"\n前30页来源:\n")
        if args.front_pages:
            f.write(f"  自定义文件: {args.front_pages}\n")
        else:
            f.write(f"  自动从项目提取（前{min(30, len(front_pages))}页）\n")

        f.write(f"\n后30页来源:\n")
        if args.back_pages:
            f.write(f"  自定义文件: {args.back_pages}\n")
        else:
            f.write(f"  自动从项目提取（后{min(30, len(back_pages))}页）\n")

    print(f"\n统计信息: {info_path}")
    print("\n完成！docx 已生成，请检查内容并按需要调整。")
    print("提醒：源代码注释率必须 <5%，可用 scripts/ai_trace_analyzer.py 核查。")


if __name__ == '__main__':
    main()
